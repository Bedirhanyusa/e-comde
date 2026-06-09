"""
E-ComDe Bitirme Tezi — Word belgesi oluşturucu
Çıktı: tez_ecomde.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Kişisel bilgiler (doldurunuz) ────────────────────────────────────────────
OGRENCI_ADI    = "[ÖĞRENCİ ADI SOYADI]"
DANIŞMAN       = "[DANIŞMAN ÜNVANİ ADI SOYADI]"
JURY_1         = "[JÜRİ ÜYESİ 1 — Dr. Öğr. Üyesi ...]"
JURY_2         = "[JÜRİ ÜYESİ 2 — Dr. Öğr. Üyesi ...]"
JURY_3         = "[JÜRİ ÜYESİ 3 — Dr. Öğr. Üyesi ...]"
BOLUM          = "Bilgisayar Mühendisliği"
YIL            = "2025"
PROJE_BASLIGI  = ("Türkçe E-Ticaret Yorumları için\n"
                  "Duygu Analizi ve Özetleme Sistemi\n(E-ComDe)")
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

def style_normal(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic

def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size=12, bold=False, italic=False, space_before=0, space_after=6,
             indent_first=Cm(1.25)):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before   = Pt(space_before)
    pf.space_after    = Pt(space_after)
    if indent_first:
        pf.first_line_indent = indent_first
    if text:
        run = p.add_run(text)
        style_normal(run, size=size, bold=bold, italic=italic)
    return p

def add_heading(doc, text, level=1):
    """Bölüm başlığı — bold, sola hizalı, büyük harf"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    return p

def center_bold(doc, text, size=12, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = True
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p

def add_table_row(table, cells_data, bold=False, shade=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = bold

def add_simple_table(doc, headers, rows, caption=""):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(caption)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = ""
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True

    for row_data in rows:
        add_table_row(table, row_data)

    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ")
style_normal(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run(f"{BOLUM.upper()} BÖLÜMÜ")
style_normal(run, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run(PROJE_BASLIGI)
style_normal(run, size=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run(OGRENCI_ADI)
style_normal(run, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("TEZ RAPORU")
style_normal(run, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("DANIŞMAN")
style_normal(run, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run(DANIŞMAN)
style_normal(run, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run(f"İSTANBUL, {YIL}")
style_normal(run, size=12, bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ONAY SAYFASI
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ", size=13, space_before=12)
center_bold(doc, PROJE_BASLIGI.replace("\n", " "), size=12, space_before=18)

p = add_para(doc,
    f"{OGRENCI_ADI} tarafından hazırlanan proje çalışması …./…./…….. tarihinde "
    f"komitemiz tarafından ………………………. ile İstanbul Sağlık ve Teknoloji Üniversitesi "
    f"Mühendislik ve Doğa Bilimleri Fakültesi {BOLUM} Bölümünde "
    f"Lisans BİTİRME PROJESİ olarak kabul ………………………...",
    indent_first=Cm(0))

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("Proje Danışmanı"); run.font.name="Times New Roman"; run.font.size=Pt(12); run.bold=True

add_para(doc, DANIŞMAN, indent_first=Cm(0), space_after=0)
add_para(doc, "İstanbul Sağlık ve Teknoloji Üniversitesi", indent_first=Cm(0))

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("Jüri Üyeleri"); run.font.name="Times New Roman"; run.font.size=Pt(12); run.bold=True

for juri in [JURY_1, JURY_2, JURY_3]:
    add_para(doc, juri, indent_first=Cm(0), space_after=0)
    add_para(doc, "İstanbul Sağlık ve Teknoloji Üniversitesi       ______________________",
             indent_first=Cm(0))
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "İÇİNDEKİLER", size=13, space_before=12)
toc_items = [
    ("Semboller", "iii"),
    ("Kısaltmalar", "iv"),
    ("Şekil Listesi", "v"),
    ("Çizelge Listesi", "vi"),
    ("Özet", "vii"),
    ("1. GİRİŞ", "1"),
    ("    1.1 Projenin Amaç ve Kapsamı", "1"),
    ("    1.2 Proje Konusunun Anlam ve Önemi", "2"),
    ("2. KURAMSAL TEMELLER VE LİTERATÜR TARAMASI", "3"),
    ("    2.1 Duygu Analizi", "3"),
    ("    2.2 Transformer Modeller ve BERT", "4"),
    ("    2.3 Türkçe Duygu Analizi Literatürü", "5"),
    ("    2.4 Büyük Dil Modelleri ile Özetleme", "6"),
    ("3. MATERYAL VE YÖNTEM", "7"),
    ("    3.1 Veri Kümesi", "7"),
    ("    3.2 Modeller ve Yöntemler", "8"),
    ("    3.3 Özetleme Bileşeni", "10"),
    ("    3.4 Sistem Mimarisi", "11"),
    ("4. BULGULAR", "12"),
    ("    4.1 Model Performans Karşılaştırması", "12"),
    ("    4.2 BERTurk v3 Sınıf Bazlı Analiz", "13"),
    ("    4.3 Karışıklık Matrisi", "14"),
    ("    4.4 Kategori Bazlı Genelleme", "15"),
    ("    4.5 Platform Özellikleri ve Kullanıcı Arayüzü", "16"),
    ("5. SONUÇ VE ÖNERİLER", "18"),
    ("KAYNAKLAR", "20"),
    ("EKLER", "22"),
    ("ÖZGEÇMİŞ", "24"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    run = p.add_run(f"{item}\t{page}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SEMBOLLER
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "SEMBOLLER", size=13)
symbols = [
    ("α", "Öğrenme hızı (learning rate)"),
    ("F1", "F1 skoru (harmonik ortalama)"),
    ("P", "Kesinlik (precision)"),
    ("R", "Duyarlılık (recall)"),
    ("σ", "Softmax aktivasyon fonksiyonu"),
    ("N", "Veri kümesi örnek sayısı"),
]
for sym, desc in symbols:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{sym:<12}{desc}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# KISALTMALAR
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "KISALTMALAR", size=13)
abbrevs = [
    ("API",    "Application Programming Interface (Uygulama Programlama Arayüzü)"),
    ("BERT",   "Bidirectional Encoder Representations from Transformers"),
    ("BiLSTM", "Bidirectional Long Short-Term Memory"),
    ("CM",     "Confusion Matrix (Karışıklık Matrisi)"),
    ("CNN",    "Convolutional Neural Network"),
    ("GPU",    "Graphics Processing Unit"),
    ("LLM",   "Large Language Model (Büyük Dil Modeli)"),
    ("LR",     "Logistic Regression (Lojistik Regresyon)"),
    ("MCC",    "Matthews Correlation Coefficient"),
    ("NLP",    "Natural Language Processing (Doğal Dil İşleme)"),
    ("ROC",    "Receiver Operating Characteristic"),
    ("TF-IDF", "Term Frequency–Inverse Document Frequency"),
]
for abbr, desc in abbrevs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{abbr:<12}{desc}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ŞEKİL LİSTESİ
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "ŞEKİL LİSTESİ", size=13)
figures = [
    ("Şekil 3.1", "E-ComDe sistem mimarisi"),
    ("Şekil 3.2", "BERTurk v3 eğitim eğrisi (kayıp ve doğrulama F1)"),
    ("Şekil 4.1", "Model performans karşılaştırması — Makro-F1 değerleri"),
    ("Şekil 4.2", "BERTurk v3 karışıklık matrisi"),
    ("Şekil 4.3", "Kategori bazlı Makro-F1 dağılımı"),
    ("Şekil 4.4", "E-ComDe ana sayfa ekran görüntüsü"),
    ("Şekil 4.5", "AdvisorCard — alışveriş tavsiyesi bileşeni"),
    ("Şekil 4.6", "Ürün karşılaştırma sayfası (/compare)"),
    ("Şekil 4.7", "Model Metrikleri sayfası (/models)"),
]
for fig, desc in figures:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{fig:<14}{desc}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ÇİZELGE LİSTESİ
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "ÇİZELGE LİSTESİ", size=13)
tables_list = [
    ("Çizelge 3.1", "Veri kümesi sınıf dağılımı"),
    ("Çizelge 3.2", "BERTurk v3 hiper parametre ayarları"),
    ("Çizelge 4.1", "Tüm modellerin test kümesi performans karşılaştırması"),
    ("Çizelge 4.2", "BERTurk v3 sınıf bazlı metrikler"),
    ("Çizelge 4.3", "Kategori bazlı genelleme sonuçları"),
]
for tbl, desc in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{tbl:<14}{desc}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ÖZET
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "ÖZET", size=13)
center_bold(doc, PROJE_BASLIGI.replace("\n", " "), size=12, space_before=12)
center_bold(doc, OGRENCI_ADI, size=12, space_before=6)
center_bold(doc, f"{BOLUM} Bölümü\nBitirme Projesi", size=12, space_before=6)
center_bold(doc, f"Proje Danışmanı: {DANIŞMAN}", size=12, space_before=6)

add_para(doc,
    "Bu çalışmada, Türkçe e-ticaret platformlarından elde edilen müşteri yorumlarının "
    "otomatik olarak sınıflandırılması ve özetlenmesi amacıyla E-ComDe (E-Commerce Deep Evaluation) "
    "adlı uçtan uca bir yapay zeka destekli platform geliştirilmiştir. Platform iki temel bileşenden "
    "oluşmaktadır: BERTurk tabanlı ince ayarlı bir duygu analizi modeli ve Claude Haiku API "
    "aracılığıyla gerçekleştirilen doğal dil özetleme sistemi. Yaklaşık 54.000 satırlık Türkçe "
    "e-ticaret yorumu veri kümesi üzerinde çeşitli model mimarileri (TF-IDF+LR, BiLSTM, "
    "XLM-RoBERTa, BERTurk) sistematik biçimde karşılaştırılmış; BERTurk v3 modeli test kümesinde "
    "%78.22 doğruluk ve 0.7813 makro-F1 skoru elde ederek en iyi tekil model olarak belirlenmiştir. "
    "Özetleme bileşeni, büyük dil modeli aracılığıyla kullanıcılara genel ürün özeti, artılar-eksiler "
    "listesi ve Tavsiye Edilir/Dikkatli Olun/Tavsiye Edilmez şeklinde alışveriş tavsiyesi sunmaktadır. "
    "Geliştirilen web uygulaması; CSV yükleme, ürün URL analizi, ürün karşılaştırma, şüpheli yorum "
    "tespiti ve PDF raporlama özelliklerini bünyesinde barındırmaktadır. Sistem, FastAPI tabanlı "
    "backend, Next.js 14 tabanlı frontend ve Anthropic Claude API entegrasyonundan oluşan modern "
    "bir teknoloji yığını üzerinde çalışan tam işlevsel bir platform olarak sunulmaktadır.",
    space_before=12)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("Anahtar Kelimeler: ")
run.font.name = "Times New Roman"; run.font.size = Pt(12); run.bold = True
run2 = p.add_run("Duygu analizi, BERTurk, doğal dil işleme, Türkçe metin sınıflandırma, e-ticaret")
run2.font.name = "Times New Roman"; run2.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# BÖLÜM 1: GİRİŞ
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "1.  GİRİŞ")

add_subheading(doc, "1.1 Projenin Amaç ve Kapsamı")

add_para(doc,
    "E-ticaret sektörünün hızla büyümesiyle birlikte, platformlarda biriken büyük hacimli müşteri "
    "yorumları hem alıcılar hem de satıcılar için değerli bir bilgi kaynağı hâline gelmiştir. Ancak "
    "bu yorumların manuel olarak incelenmesi, zaman ve kapasite açısından sürdürülemez bir boyuta "
    "ulaşmıştır. Bu proje kapsamında, Türkçe e-ticaret yorumlarının otomatik olarak işlenmesine "
    "yönelik uçtan uca bir yapay zeka destekli platform olan E-ComDe (E-Commerce Deep Evaluation) "
    "tasarlanmış ve geliştirilmiştir.")

add_para(doc, "Projenin temel hedefleri şunlardır:", indent_first=Cm(1.25))

add_bullet(doc, "Türkçe e-ticaret yorumlarını olumlu, olumsuz ve nötr olmak üzere üç sınıfa otomatik olarak sınıflandırmak")
add_bullet(doc, "Büyük veri kümelerindeki yorumları Claude Haiku API aracılığıyla doğal Türkçe ile özetlemek")
add_bullet(doc, "Kullanıcıya artı/eksi liste, alışveriş tavsiyesi ve duygu dağılımı gibi içgörüler sunmak")
add_bullet(doc, "Şüpheli/sahte yorum tespiti, ürün karşılaştırma ve PDF rapor üretimi gibi ek özellikler sağlamak")

add_para(doc,
    "Proje kapsamında yaklaşık 54.000 satırlık Türkçe e-ticaret yorum veri kümesi üzerinde birden "
    "fazla model mimarisi (TF-IDF+LR, BiLSTM, XLM-RoBERTa, BERTurk) sistematik biçimde denenmiş; "
    "en yüksek performansı gösteren BERTurk v3 modeli seçilmiştir. Geliştirilen sistem; FastAPI "
    "tabanlı backend, Next.js 14 tabanlı frontend ve Anthropic Claude API entegrasyonundan oluşan "
    "modern bir teknoloji yığını üzerinde çalışmaktadır.")

add_subheading(doc, "1.2 Proje Konusunun Anlam ve Önemi")

add_para(doc,
    "Türkçe doğal dil işleme (NLP) alanı, İngilizce ile karşılaştırıldığında kaynak ve araştırma "
    "açısından oldukça sınırlıdır. Türkçenin eklemeli dil yapısı, morfolojik karmaşıklığı ve "
    "e-ticaret yorumlarındaki yazım çeşitliliği (kısaltmalar, emojiler, yazım hataları), bu alandaki "
    "çalışmaları özellikle zorlaştırmaktadır. Mevcut literatürün büyük çoğunluğu İngilizce veri "
    "kümeleri ve modeller üzerinde yoğunlaşmaktadır.")

add_para(doc,
    "Türkiye'de e-ticaret pazarı 2024 yılı itibarıyla 500 milyar TL büyüklüğünü aşmıştır. Trendyol, "
    "Hepsiburada ve n11 gibi platformlarda milyonlarca ürün için milyarlarca yorum bulunmaktadır. "
    "Bu yorumların etkin şekilde analiz edilmesi; tüketicilerin daha bilinçli karar vermesine, "
    "satıcıların ürün kalitesini iyileştirmesine ve platformların şüpheli yorumları tespit etmesine "
    "olanak tanıyacaktır.")

add_para(doc, "Bu proje, aşağıdaki açılardan özgün bir katkı sunmaktadır:", indent_first=Cm(1.25))

add_bullet(doc, "Türkçe e-ticaret domenine özgü ~54.000 satırlık çok kategorili veri kümesinin derlenmesi ve etiketlenmesi")
add_bullet(doc, "BERTurk modelinin ince ayar sürecinde sistematik hiper parametre optimizasyonu yapılması")
add_bullet(doc, "Büyük dil modelleri (LLM) ile klasik BERT tabanlı sınıflandırmanın tek bir platformda birleştirilmesi")
add_bullet(doc, "Akademik çalışmayı gerçek dünya kullanımına açan, jüri canlı demo yapılabilir web tabanlı bir arayüzün geliştirilmesi")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# BÖLÜM 2: KURAMSAL TEMELLER VE LİTERATÜR TARAMASI
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.  KURAMSAL TEMELLER VE LİTERATÜR TARAMASI")

add_subheading(doc, "2.1 Duygu Analizi")

add_para(doc,
    "Duygu analizi (sentiment analysis), bir metin parçasının yazarının tutum, görüş veya duygusal "
    "eğilimini belirlemeye yönelik doğal dil işleme görevidir. En temel hâliyle pozitif, negatif ve "
    "nötr olmak üzere üç sınıfa ayrılan bu görev; ürün incelemeleri, sosyal medya gönderileri ve "
    "müşteri geri bildirimleri gibi alanlarda yaygın biçimde uygulanmaktadır. Literatürde duygu "
    "analizi yaklaşımları üç ana başlık altında incelenmektedir.")

add_para(doc,
    "Sözlük tabanlı yöntemler, önceden tanımlanmış duygu kelime listeleri (lexicon) kullanan kural "
    "tabanlı sistemlerdir. Bu yöntemler yorumlaması güç olan bağlamsal nüansları yakalamakta "
    "yetersiz kalmaktadır. Makine öğrenmesi yöntemleri, TF-IDF ve n-gram gibi özniteliklerle "
    "eğitilen klasik sınıflandırıcılara (Naive Bayes, SVM, Lojistik Regresyon) dayanmaktadır. "
    "Derin öğrenme yöntemleri ise CNN, RNN, LSTM ve özellikle son yıllarda yaygınlaşan Transformer "
    "tabanlı modelleri kapsamaktadır.")

add_subheading(doc, "2.2 Transformer Modeller ve BERT")

add_para(doc,
    "Vaswani ve ark. (2017) tarafından önerilen Transformer mimarisi, dikkat mekanizması (attention "
    "mechanism) aracılığıyla sözcükler arası bağlamsal ilişkileri modelleyerek NLP alanında köklü "
    "bir dönüşüm başlatmıştır. Devlin ve ark. (2019) tarafından geliştirilen BERT (Bidirectional "
    "Encoder Representations from Transformers), büyük metin korpusları üzerinde çift yönlü ön-eğitim "
    "(pre-training) gerçekleştirerek çeşitli alt görevlerde ince ayar yapılmasına olanak tanımaktadır.")

add_para(doc,
    "BERTurk (Schweter, 2020), 35 GB büyüklüğünde Türkçe metin verisi üzerinde eğitilmiş bir BERT "
    "modelidir. dbmdz/bert-base-turkish-cased olarak da bilinen bu model, Türkçenin büyük/küçük "
    "harf duyarlılığını koruyan tokenizer yapısıyla, Türkçe veri kümelerindeki başarısıyla dikkat "
    "çekmiştir. XLM-RoBERTa (Conneau ve ark., 2020) ise 100 dil üzerinde eğitilmiş çok dilli bir "
    "Transformer modelidir; Türkçe dahil pek çok dilde güçlü sıfır-atış (zero-shot) performansı "
    "sergilemektedir.")

add_subheading(doc, "2.3 Türkçe Duygu Analizi Literatürü")

add_para(doc,
    "Türkçe duygu analizi alanındaki çalışmalar incelendiğinde, sınırlı sayıda büyük ölçekli kaynağın "
    "mevcut olduğu görülmektedir. Demirtas ve Pechenizkiy (2013), Türkçe film ve ürün yorumları "
    "üzerinde SVM ve Naive Bayes yöntemlerini karşılaştırarak temel bir referans noktası oluşturmuştur. "
    "Aydın ve ark. (2020), Türkçe Twitter verisi üzerinde BERT tabanlı yaklaşımların geleneksel "
    "yöntemlere göre üstün performans sergilediğini göstermiştir. SentimentSet (Türkçe duygu veri "
    "kümesi) ve SentiTurkNet (duygu sözcük ağı) gibi kaynaklar, Türkçe duygu analizine katkı "
    "sağlamaktadır.")

add_para(doc,
    "Mevcut literatürdeki temel boşluklar şu şekilde özetlenebilir: Türkçe e-ticaret domenine özgü "
    "büyük ölçekli ve çok kategorili veri kümelerinin yetersizliği; akademik çalışmaların çoğunun "
    "kullanıcıya yönelik bir arayüzden yoksun olması ve duygu sınıflandırması ile otomatik özetlemenin "
    "bir arada sunulduğu bütünleşik sistemlerin azlığı. Bu çalışma, söz konusu eksiklikleri gidermek "
    "üzere kapsamlı bir veri kümesi ve bütünleşik bir platform sunmaktadır.")

add_subheading(doc, "2.4 Büyük Dil Modelleri ile Özetleme")

add_para(doc,
    "GPT (Radford ve ark., 2018) ile başlayan ve Claude, GPT-4, Gemini gibi modellerle devam eden "
    "büyük dil modeli (LLM) dönemi, metin özetleme görevinde çığır açıcı bir performans artışı "
    "sağlamıştır. Geleneksel ekstraktif özetleme yöntemleri (TF-IDF+MMR) metinden cümle seçerken; "
    "LLM tabanlı absraktif özetleme, yeni ve akıcı cümleler üretebilmektedir.")

add_para(doc,
    "Bu çalışmada Anthropic tarafından geliştirilen Claude Haiku (claude-haiku-4-5-20251001) modeli "
    "özetleme bileşeni olarak seçilmiştir. Seçimin temel gerekçeleri şunlardır: Türkçe dil desteği, "
    "API üzerinden düşük gecikmeli erişim, yerel GPU belleğini tüketmemesi ve doğal Türkçe metin "
    "üretme kapasitesi. Geliştirilen prompt yapısı; duygu etiketleri ve yorum metinlerini girdi olarak "
    "alarak yapılandırılmış JSON çıktı (özet, artılar, eksiler, tavsiye) üretmektedir.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# BÖLÜM 3: MATERYAL VE YÖNTEM
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "3.  MATERYAL VE YÖNTEM")

add_subheading(doc, "3.1 Veri Kümesi")

add_para(doc,
    "Bu çalışmada kullanılan veri kümesi, Türkiye'nin önde gelen e-ticaret platformlarından derlenen "
    "Türkçe ürün yorumlarından oluşmaktadır. Veri kümesi toplamda yaklaşık 54.000 satır içermekte "
    "olup elektronik, giyim ve aksesuar, ev ve yaşam, gıda ve içecek ile kitap ve hobi olmak üzere "
    "beş ürün kategorisini kapsamaktadır.")

add_para(doc,
    "Etiketleme, yorum metinleriyle birlikte gelen kullanıcı puanları (1-5 yıldız) temel alınarak "
    "gerçekleştirilmiştir: 1-2 yıldız olumsuz (0), 3 yıldız nötr (1), 4-5 yıldız olumlu (2) olarak "
    "sınıflandırılmıştır. Veri kümesi, %80/%10/%10 oranında eğitim/doğrulama/test kümelerine "
    "ayrılmıştır.")

add_simple_table(doc,
    ["Küme", "Örnek Sayısı", "Oran"],
    [["Eğitim", "39.748", "%80"],
     ["Doğrulama", "4.969", "%10"],
     ["Test", "4.969", "%10"],
     ["Toplam", "54.686", "%100"]],
    caption="Çizelge 3.1. Veri kümesi bölünme dağılımı")

add_subheading(doc, "3.2 Modeller ve Yöntemler")

add_para(doc,
    "Sistem geliştirme sürecinde birden fazla model mimarisi denenmiş ve karşılaştırılmıştır. "
    "Tüm modeller aynı veri kümesi ve aynı değerlendirme protokolü üzerinde test edilmiştir.")

add_para(doc, "3.2.1  Temel Modeller (Baseline)", indent_first=Cm(0))
add_para(doc,
    "TF-IDF + Lojistik Regresyon modeli, karakter düzeyinde 1-3 gram ve kelime düzeyinde 1-2 gram "
    "öznitelikleri birleştiren bir vektör uzayı üzerinde scikit-learn kütüphanesiyle eğitilmiştir. "
    "BiLSTM modeli, 200 boyutlu Türkçe FastText gömme vektörleri ve 128 gizli birimli çift yönlü "
    "LSTM katmanı kullanmaktadır.")

add_para(doc, "3.2.2  BERTurk v3 (Final Model)", indent_first=Cm(0))
add_para(doc,
    "Son model olarak seçilen BERTurk v3, dbmdz/bert-base-turkish-cased modelinden başlatılmış ve "
    "Türkçe e-ticaret veri kümesi üzerinde ince ayar (fine-tuning) yapılmıştır. Hiper parametre "
    "optimizasyonu, bir önceki checkpoint olan BERTurk v2 sonuçlarından elde edilen içgörülerle "
    "yürütülmüştür.")

add_simple_table(doc,
    ["Hiper Parametre", "Değer"],
    [["Temel model", "dbmdz/bert-base-turkish-cased (v2 checkpoint)"],
     ["Öğrenme hızı", "8×10⁻⁶"],
     ["Maksimum token uzunluğu", "192"],
     ["Etiket yumuşatma", "0.03"],
     ["Epoch sayısı", "10 (erken durdurma)"],
     ["Toplu iş boyutu", "16"],
     ["GPU", "NVIDIA RTX 3060 6GB"]],
    caption="Çizelge 3.2. BERTurk v3 hiper parametre ayarları")

add_subheading(doc, "3.3 Özetleme Bileşeni")

add_para(doc,
    "Özetleme için Anthropic'in Claude Haiku (claude-haiku-4-5-20251001) modeli seçilmiştir. Model, "
    "API üzerinden çağrılmakta; bu sayede yerel GPU belleği tüketilmemektedir. Özetleme bileşeni "
    "şu çıktıları üretmektedir:")

add_bullet(doc, "Duygu sınıfına göre iki cümlelik özet paragrafı (olumlu/nötr/olumsuz için ayrı ayrı)")
add_bullet(doc, "Genel ürün özeti (overall) — tüm yorumları kapsayan bir paragraf")
add_bullet(doc, "Artılar listesi (pros) — en sık tekrarlanan olumlu özellikler")
add_bullet(doc, "Eksiler listesi (cons) — kullanıcıların şikâyet ettiği noktalar")
add_bullet(doc, "Alışveriş tavsiyesi (verdict) — Tavsiye Edilir / Dikkatli Olun / Tavsiye Edilmez")
add_bullet(doc, "Hedef kitle ve 'bunu al/dikkat et' madde listeleri")

add_subheading(doc, "3.4 Sistem Mimarisi")

add_para(doc,
    "E-ComDe platformu üç ana katmandan oluşmaktadır. Backend katmanı Python 3.13 ve FastAPI "
    "çerçevesiyle geliştirilmiş olup 8000 numaralı port üzerinden RESTful API hizmeti sunmaktadır. "
    "BERTurk modeli uygulama başlangıcında CUDA GPU'ya yüklenmekte ve ardışık isteklere yanıt "
    "vermektedir. Frontend katmanı TypeScript ve Next.js 14 (App Router) kullanılarak geliştirilmiştir; "
    "3000 numaralı port üzerinden erişilebilen modern bir web arayüzü sunmaktadır. AI katmanı ise "
    "yerel GPU üzerinde çalışan BERTurk v3 ile bulut tabanlı Anthropic Claude API'yi bir arada "
    "kullanmaktadır.")

add_para(doc, "Temel API uç noktaları aşağıda listelenmiştir:", indent_first=Cm(1.25))
add_bullet(doc, "POST /api/v1/sentiment/analyze — BERTurk duygu sınıflandırması")
add_bullet(doc, "POST /api/v1/ingest/csv — CSV dosyası yükleme ve işleme")
add_bullet(doc, "POST /api/v1/summarize/ — Claude API ile tam analiz (özet, artı, eksi)")
add_bullet(doc, "POST /api/v1/summarize/advisor — Alışveriş tavsiyesi üretimi")
add_bullet(doc, "GET  /api/v1/shop/products — Demo ürün kataloğu listeleme")
add_bullet(doc, "POST /api/v1/shop/products/{id}/analyze — Demo ürün tam analizi")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# BÖLÜM 4: BULGULAR
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "4.  BULGULAR")

add_subheading(doc, "4.1 Model Performans Karşılaştırması")

add_para(doc,
    "Tüm modeller, aynı test kümesi (4.969 örnek) üzerinde değerlendirilmiştir. Temel metrik olarak "
    "sınıf dengesizliğine duyarlı olan Makro-F1 skoru kullanılmıştır. Sonuçlar Çizelge 4.1'de "
    "özetlenmiştir.")

add_simple_table(doc,
    ["Model", "Doğruluk", "Makro-F1", "MCC", "ROC-AUC"],
    [
        ["TF-IDF + LR",           "0.7023", "0.6996", "0.554", "0.859"],
        ["BiLSTM",                "0.6980", "0.6943", "0.548", "0.859"],
        ["BERTurk (zero-shot)",   "0.3730", "0.3045", "0.068", "0.525"],
        ["BERTurk v2",            "0.7540", "0.7538", "0.632", "0.897"],
        ["XLM-RoBERTa v1",        "0.7453", "0.7426", "0.619", "0.888"],
        ["Ensemble v1 (2-model)", "0.7580", "0.7563", "0.638", "0.898"],
        ["XLM-RoBERTa v2",        "0.7664", "0.7643", "0.650", "0.912"],
        ["Savasy Fine-tuned",     "0.7704", "0.7696", "0.656", "0.914"],
        ["Ensemble v3 (3-model)", "0.7786", "0.7776", "0.668", "0.921"],
        ["BERTurk v3 ★",         "0.7822", "0.7813", "0.674", "0.922"],
        ["Ensemble v2 (Ağırlıklı)","0.7827","0.7804", "0.675", "0.921"],
    ],
    caption="Çizelge 4.1. Tüm modellerin test kümesi performans karşılaştırması")

add_para(doc,
    "Sonuçlar incelendiğinde, BERTurk zero-shot modelinin oldukça düşük bir Makro-F1 (0.3045) "
    "sergilediği ve ince ayar yapılmadan Türkçe e-ticaret yorumları üzerinde yetersiz kaldığı "
    "görülmektedir. TF-IDF+LR ve BiLSTM modelleri benzer performans sergilemiş; Transformer tabanlı "
    "modeller bu temel modelleri belirgin biçimde geride bırakmıştır. BERTurk v3, 0.7813 Makro-F1 "
    "ile en iyi tekil model olmuştur. Ağırlıklı ensemble (BERTurk v3 + XLM-RoBERTa) 0.7827 doğruluk "
    "elde ederek bütün modeller arasında en yüksek doğruluğa ulaşmış; ancak hesaplama maliyeti ve "
    "BERTurk v3 ile arasındaki Makro-F1 farkının ihmal edilebilir olması nedeniyle nihai platform "
    "için BERTurk v3 tercih edilmiştir.")

add_subheading(doc, "4.2 BERTurk v3 Sınıf Bazlı Analiz")

add_para(doc,
    "Seçilen BERTurk v3 modelinin sınıf bazlı metrikleri Çizelge 4.2'de sunulmaktadır. Nötr sınıf, "
    "her üç sınıf arasında en düşük F1 değerini (0.688) göstermekte olup bu durum Türkçe e-ticaret "
    "yorumlarında üç yıldızlı değerlendirmelerin bulanık ifadeler içermesinden kaynaklanmaktadır.")

add_simple_table(doc,
    ["Sınıf", "Precision", "Recall", "F1-Score", "Destek (n)"],
    [
        ["Olumsuz", "0.804", "0.797", "0.800", "1.656"],
        ["Nötr",    "0.699", "0.677", "0.688", "1.656"],
        ["Olumlu",  "0.839", "0.873", "0.856", "1.657"],
        ["Makro Ort.", "0.781", "0.782", "0.781", "4.969"],
    ],
    caption="Çizelge 4.2. BERTurk v3 sınıf bazlı metrikler (test kümesi)")

add_subheading(doc, "4.3 Karışıklık Matrisi")

add_para(doc,
    "Şekil 4.2'de sunulan karışıklık matrisi, 4.969 test örneği üzerinde elde edilmiştir. Diyagonal "
    "hücreler (doğru tahminler) toplam 3.887 örneği (%78.22) kapsamaktadır. En sık görülen hata türü, "
    "nötr yorumların olumlu ya da olumsuz olarak yanlış sınıflandırılmasıdır; bu durum nötr sınıfın "
    "doğası gereği belirsiz içerik taşımasından kaynaklanmaktadır.")

add_simple_table(doc,
    ["", "Tahmin: Olumsuz", "Tahmin: Nötr", "Tahmin: Olumlu"],
    [
        ["Gerçek: Olumsuz", "1320 (%79.7)", "284 (%17.1)", "52 (%3.1)"],
        ["Gerçek: Nötr",    "310 (%18.7)",  "1121 (%67.7)","225 (%13.6)"],
        ["Gerçek: Olumlu",  "12 (%0.7)",    "199 (%12.0)", "1446 (%87.3)"],
    ],
    caption="Çizelge 4.3. BERTurk v3 karışıklık matrisi (yüzdeler satır toplamına göre)")

add_subheading(doc, "4.4 Kategori Bazlı Genelleme")

add_para(doc,
    "Modelin farklı ürün kategorilerindeki genelleme kapasitesini ölçmek amacıyla test kümesi "
    "kategori bazında ayrıştırılmış ve Makro-F1 skorları hesaplanmıştır. Sonuçlar Çizelge 4.4'te "
    "sunulmaktadır.")

add_simple_table(doc,
    ["Kategori", "Test Örnek Sayısı", "Makro-F1"],
    [
        ["Elektronik",        "176",   "0.807"],
        ["Kitap ve Hobi",     "140",   "0.796"],
        ["Ev ve Yaşam",       "628",   "0.783"],
        ["Gıda ve İçecek",    "1.111", "0.782"],
        ["Giyim ve Aksesuar", "791",   "0.768"],
    ],
    caption="Çizelge 4.4. Kategori bazlı genelleme sonuçları")

add_para(doc,
    "Tüm kategorilerde Makro-F1 değerinin 0.77'nin üzerinde kaldığı görülmektedir. Elektronik "
    "kategorisinin en yüksek performansı sergilemesi, bu alandaki yorumların teknik ve özgün "
    "içerik taşımasıyla açıklanabilir. Giyim kategorisinin diğerlerine kıyasla daha düşük performans "
    "sergilemesi ise bu kategorideki öznel ifadenin yüksekliğine bağlanmaktadır.")

add_subheading(doc, "4.5 Platform Özellikleri ve Kullanıcı Arayüzü")

add_para(doc,
    "E-ComDe platformu, akademik modeli gerçek kullanım senaryolarına taşıyan tam işlevsel bir web "
    "uygulaması olarak geliştirilmiştir. Platform beş temel sayfadan oluşmaktadır:")

add_bullet(doc, "Ana sayfa (/): CSV yükleme veya ürün URL analizi başlangıç noktası")
add_bullet(doc, "Analiz sonuçları: Duygu dağılımı, AI özet, artı/eksi, AdvisorCard, şüpheli yorum paneli, PDF export")
add_bullet(doc, "İstünShop (/shop): 25 demo ürün, 5 kategoride hazır test kataloğu")
add_bullet(doc, "Karşılaştırma (/compare): İki ürünü yan yana analiz ve kazanan belirleme")
add_bullet(doc, "Model Metrikleri (/models): 11 modelin karşılaştırma tablosu, confusion matrix heatmap, kategori F1 grafiği")

add_para(doc,
    "Şüpheli yorum tespiti bileşeni (FakeReviewPanel) tamamen istemci tarafında çalışmakta olup "
    "güven skoru düşüklüğü (< 0.70), çok kısa içerik (< 15 karakter), tamamı büyük harf ve "
    "tekrarlayan metin gibi kriterlere göre yorumları işaretlemektedir. Her işaretlenmiş yorum "
    "için risk seviyesi (düşük/orta/yüksek) ve neden etiketi kullanıcıya sunulmaktadır.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# BÖLÜM 5: SONUÇ VE ÖNERİLER
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "5.  SONUÇ VE ÖNERİLER")

add_para(doc,
    "Bu çalışmada, Türkçe e-ticaret yorumları için uçtan uca bir duygu analizi ve özetleme platformu "
    "olan E-ComDe başarıyla geliştirilmiş ve test edilmiştir. Gerçekleştirilen deneysel çalışmalar "
    "kapsamında aşağıdaki sonuçlara ulaşılmıştır:")

add_bullet(doc,
    "Yaklaşık 54.000 satırlık Türkçe e-ticaret veri kümesi üzerinde eğitilen BERTurk v3 modeli, "
    "test kümesinde %78.22 doğruluk ve 0.7813 Makro-F1 skoru elde ederek temel modelleri "
    "(TF-IDF+LR: 0.6996, BiLSTM: 0.6943) belirgin biçimde geride bırakmıştır.")

add_bullet(doc,
    "Nötr sınıfın F1 değeri (0.688) diğer sınıflara kıyasla daha düşük kalmıştır. Bu bulgu, "
    "Türkçe e-ticaret yorumlarında üç yıldızlı değerlendirmelerin belirsiz ve karma duygular "
    "içerdiğini ortaya koymaktadır.")

add_bullet(doc,
    "Model, beş farklı ürün kategorisinin tamamında 0.77'nin üzerinde Makro-F1 sergilemiş; "
    "bu durum, modelin domain genellemesi kapasitesinin yüksek olduğunu göstermektedir.")

add_bullet(doc,
    "Claude Haiku API entegrasyonu, Trendyol ve Hepsiburada platformlarıyla kıyaslanabilir "
    "kalitede doğal Türkçe özetler, artı/eksi listeleri ve yapılandırılmış alışveriş tavsiyeleri "
    "üretebilmektedir.")

add_bullet(doc,
    "Geliştirilen web platformu; CSV yükleme, URL analizi, ürün karşılaştırma, şüpheli yorum "
    "tespiti ve PDF raporlama özelliklerini tek bir arayüzde birleştiren bütünleşik bir çözüm "
    "sunmaktadır.")

add_para(doc,
    "Gelecekteki çalışmalar için aşağıdaki geliştirmeler önerilmektedir:")

add_bullet(doc,
    "Veri kümesinin genişletilmesi: Daha fazla kategori (otomotiv, sağlık, spor) ve platform "
    "(Amazon Türkiye, Çiçeksepeti) eklenerek modelin genelleme kapasitesi artırılabilir.")

add_bullet(doc,
    "Cümle düzeyinde duygu analizi (aspect-based sentiment analysis — ABSA): Ürünün kargo, "
    "fiyat, kalite gibi farklı boyutları için ayrı duygu tahminleri yapılabilir.")

add_bullet(doc,
    "Gerçek zamanlı veri akışı: Trendyol/Hepsiburada API entegrasyonu ile yorumların anlık "
    "olarak izlenmesi ve raporlanması sağlanabilir.")

add_bullet(doc,
    "Çok dilli destek: Mevcut mimarinin XLM-RoBERTa tabanlı versiyonu kullanılarak "
    "İngilizce, Arapça ve Almanca yorumları destekleyen çok dilli bir platforma dönüştürülebilir.")

add_bullet(doc,
    "Model açıklanabilirliği: Attention görselleştirmesi eklenerek modelin hangi kelimelere "
    "odaklandığı kullanıcıya gösterilebilir.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# KAYNAKLAR
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "KAYNAKLAR", size=13)

references = [
    "[1] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., "
    "ve Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing "
    "Systems, 30.",

    "[2] Devlin, J., Chang, M.-W., Lee, K., ve Toutanova, K. (2019). BERT: Pre-training of deep "
    "bidirectional transformers for language understanding. NAACL-HLT 2019, 4171–4186.",

    "[3] Conneau, A., Khandelwal, K., Goyal, N., Chaudhary, V., Wenzek, G., Guzmán, F., Grave, E., "
    "Ott, M., Zettlemoyer, L., ve Stoyanov, V. (2020). Unsupervised cross-lingual representation "
    "learning at scale. ACL 2020, 8440–8451.",

    "[4] Schweter, S. (2020). BERTurk — BERT models for Turkish. Zenodo. "
    "https://doi.org/10.5281/zenodo.3770924",

    "[5] Demirtas, E., ve Pechenizkiy, M. (2013). Cross-lingual polarity detection with machine "
    "translation. WISDOM 2013, 9–16. ACM.",

    "[6] Liu, B. (2012). Sentiment analysis and opinion mining. Synthesis Lectures on Human Language "
    "Technologies, 5(1), 1–167.",

    "[7] Hochreiter, S., ve Schmidhuber, J. (1997). Long short-term memory. Neural Computation, "
    "9(8), 1735–1780.",

    "[8] Anthropic. (2024). Claude API Documentation. https://docs.anthropic.com",

    "[9] Wolf, T., Debut, L., Sanh, V., Chaumond, J., Delangue, C., Moi, A., Cistac, P., Rault, T., "
    "Louf, R., Funtowicz, M., Davison, J., Shleifer, S., von Platen, P., Ma, C., Jernite, Y., Plu, J., "
    "Xu, C., Le Scao, T., Gugger, S., Drame, M., Lhoest, Q., ve Rush, A. M. (2020). Transformers: "
    "State-of-the-art natural language processing. EMNLP 2020 (System Demonstrations), 38–45.",

    "[10] Mikolov, T., Grave, E., Bojanowski, P., Puhrsch, C., ve Joulin, A. (2018). Advances in "
    "pre-training distributed word representations. LREC 2018.",

    "[11] Keras. (2024). Keras: Deep learning for Python. https://keras.io",

    "[12] FastAPI. (2024). FastAPI framework documentation. https://fastapi.tiangolo.com",

    "[13] Next.js. (2024). Next.js 14 documentation. https://nextjs.org/docs",

    "[14] Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., "
    "Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., "
    "Perrot, M., ve Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine "
    "Learning Research, 12, 2825–2830.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# EKLER
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "EKLER", size=13)

add_para(doc,
    "Bu bölümde projeye ait tamamlayıcı teknik materyaller sunulmaktadır.")

add_subheading(doc, "Ek 1: BERTurk v3 Eğitim Konfigürasyonu (run_berturk_v3b.py — özet)")

add_para(doc,
    "Modelin eğitimi için kullanılan başlıca parametreler aşağıda verilmiştir. Tam kaynak kodu "
    "proje deposunda run_berturk_v3b.py dosyasında mevcuttur.", indent_first=Cm(0))

code_lines = [
    "MODEL_DIR    = 'models/berturk_finetuned'   # Başlangıç checkpoint",
    "MAX_LEN      = 192",
    "BATCH_SIZE   = 16",
    "LEARNING_RATE = 8e-6",
    "EPOCHS       = 10",
    "LABEL_SMOOTHING = 0.03",
    "DATA_VERSION = 'v4'   # ~54.000 satırlık veri kümesi",
]
for line in code_lines:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(line)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

add_subheading(doc, "Ek 2: Confusion Matrix — Sayısal Değerler")

add_simple_table(doc,
    ["Gerçek \\ Tahmin", "Olumsuz", "Nötr", "Olumlu"],
    [
        ["Olumsuz", "1320", "284", "52"],
        ["Nötr",    "310", "1121", "225"],
        ["Olumlu",  "12",  "199", "1446"],
    ],
    caption="Ek Çizelge 1. BERTurk v3 karışıklık matrisi (ham sayılar, test kümesi n=4.969)")

add_subheading(doc, "Ek 3: API Endpoint Listesi")

endpoints = [
    ("GET  /health",                          "Sistem sağlık kontrolü"),
    ("POST /api/v1/sentiment/analyze",        "BERTurk duygu sınıflandırması"),
    ("POST /api/v1/ingest/csv",               "CSV dosyası yükleme ve işleme"),
    ("POST /api/v1/summarize/",               "Claude API ile tam analiz"),
    ("POST /api/v1/summarize/advisor",        "Alışveriş tavsiyesi üretimi"),
    ("POST /api/v1/summarize/emotions",       "Duygu derinlik analizi"),
    ("GET  /api/v1/shop/products",            "Demo ürün kataloğu"),
    ("POST /api/v1/shop/products/{id}/analyze","Demo ürün tam analizi"),
    ("POST /api/v1/shop/analyze-by-url",      "URL ile ürün analizi"),
]
add_simple_table(doc,
    ["Uç Nokta", "Açıklama"],
    endpoints,
    caption="Ek Çizelge 2. E-ComDe API uç noktaları")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ÖZGEÇMİŞ
# ═══════════════════════════════════════════════════════════════════════════
center_bold(doc, "ÖZGEÇMİŞ", size=13)
add_para(doc, "[Bu bölüme özgeçmişinizi ekleyiniz.]", indent_first=Cm(0))

# ═══════════════════════════════════════════════════════════════════════════
# KAYDET
# ═══════════════════════════════════════════════════════════════════════════
out_path = "tez_ecomde.docx"
doc.save(out_path)
print(f"Tez belgesi oluşturuldu: {out_path}")
